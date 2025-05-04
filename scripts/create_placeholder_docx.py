from docx import Document
from docx.shared import Inches
import os

def create_placeholder_docx(filepath):
    doc = Document()
    doc.add_heading('Medicaid Insurance Rates', 0)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ICD Code'
    hdr_cells[1].text = 'Disease'
    hdr_cells[2].text = 'Default Rate'
    # Add a sample row
    row_cells = table.add_row().cells
    row_cells[0].text = 'A00'
    row_cells[1].text = 'Cholera'
    row_cells[2].text = '100'
    doc.save(filepath)

if __name__ == "__main__":
    target_path = os.path.join(os.path.dirname(__file__), '../data/Medicaid_Insurance_Rates.docx')
    create_placeholder_docx(target_path)
    print(f"Placeholder DOCX created at: {target_path}")
