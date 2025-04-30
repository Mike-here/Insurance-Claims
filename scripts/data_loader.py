import os
import pandas as pd
from docx import Document

from io import BytesIO
import pandas as pd
from docx import Document
import os
import numpy as np

def docx_table_to_df(docx_file):
    # Accept file path or file-like object
    if isinstance(docx_file, str):
        doc = Document(docx_file)
    else:
        doc = Document(docx_file)
    table = doc.tables[0]
    data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    df = pd.DataFrame(data[1:], columns=data[0])
    return df

def csv_to_df(csv_file):
    return pd.read_csv(csv_file)

def load_default_data():
    """Load default files from data/ directory."""
    # Get absolute path to script directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    data_dir = os.path.join(script_dir, '..', 'data')
    insurance_rates_dir = os.path.join(data_dir, 'insurance_rates')
    
    # Ensure we're using absolute paths
    data_dir = os.path.abspath(data_dir)
    insurance_rates_dir = os.path.abspath(insurance_rates_dir)
    
    print(f"Data directory: {data_dir}")
    print(f"Insurance rates directory: {insurance_rates_dir}")
    
    # Load doctor and patient data
    doctor_file_path = os.path.join(data_dir, 'Doctor_Charges.docx')
    print(f"Loading doctor file from: {doctor_file_path}")
    doctor_df = docx_table_to_df(doctor_file_path)
    
    patient_file_path = os.path.join(data_dir, 'Patient_Disease_Assignments.docx')
    print(f"Loading patient file from: {patient_file_path}")
    patient_df = docx_table_to_df(patient_file_path)
    
    # List of valid insurance companies
    valid_insurances = ['Aetna', 'BlueCross', 'Medicare', 'UnitedHealthCare', 'Medicaid']
    
    # Get list of all .docx files in the directory
    insurance_files = [f for f in os.listdir(insurance_rates_dir) 
                      if f.endswith('.docx') and not f.startswith('~$')]
    
    print(f"Found insurance files: {insurance_files}")
    
    insurance_rates = {}
    
    for insurance_file in insurance_files:
        try:
            # Get the insurance name from the filename (without extension)
            insurance_name = os.path.splitext(insurance_file)[0]
            
            # Clean up the insurance name:
            # 1. Convert to lowercase
            # 2. Remove '_rates' or '_insurance_rates'
            # 3. Replace spaces with underscores
            insurance_name = insurance_name.lower()
            if insurance_name.endswith('_rates'):
                insurance_name = insurance_name[:-6]
            elif insurance_name.endswith('_insurance_rates'):
                insurance_name = insurance_name[:-15]
            insurance_name = insurance_name.replace(' ', '_')
            
            # Map the cleaned name back to our standard names
            name_mapping = {
                'aetna': 'Aetna',
                'bluecross': 'BlueCross',
                'medicare': 'Medicare',
                'unitedhealthcare': 'UnitedHealthCare',
                'medicaid': 'Medicaid'
            }
            
            if insurance_name in name_mapping:
                insurance_name = name_mapping[insurance_name]
            else:
                print(f"Warning: Unrecognized insurance name: {insurance_name}")
                continue
            
            # Skip if it's not one of our valid insurance companies
            if insurance_name not in valid_insurances:
                print(f"Skipping invalid insurance file: {insurance_file}")
                continue
            
            insurance_path = os.path.join(insurance_rates_dir, insurance_file)
            print(f"\nLoading insurance file: {insurance_path}")
            
            # Try to open and read the file
            try:
                insurance_df = docx_table_to_df(insurance_path)
            except Exception as e:
                print(f"Error reading {insurance_file}: {str(e)}")
                continue
            
            # Print column names for debugging
            print(f"Columns in {insurance_file}: {insurance_df.columns.tolist()}")
            
            # Check if required columns exist
            required_columns = ['Disease Name', 'Insurance Rate']
            missing_columns = [col for col in required_columns if col not in insurance_df.columns]
            if missing_columns:
                print(f"Warning: Missing columns in {insurance_file}: {missing_columns}")
                continue
            
            # Create dictionary of coverage amounts for this insurance
            insurance_rates[insurance_name] = {
                row['Disease Name']: float(row['Insurance Rate'])
                for _, row in insurance_df.iterrows()
            }
            
            # Print loaded rates for debugging
            print(f"Loaded rates for {insurance_name}:")
            for disease, rate in insurance_rates[insurance_name].items():
                print(f"  {disease}: ${rate:.2f}")
                
        except Exception as e:
            print(f"Error processing {insurance_file}: {str(e)}")
            continue
    
    billing_df = prepare_billing_summary(doctor_df, insurance_rates, patient_df)
    
    return doctor_df, insurance_rates, patient_df, billing_df

def prepare_billing_summary(doctor_df, insurance_rates, patient_df):
    """
    Prepare a billing summary that shows:
    1. Total charges for each doctor
    2. Insurance coverage for each patient
    3. Patient's out-of-pocket costs
    """
    billing_summary = []
    
    # Group patients by doctor
    for doctor in patient_df['Assigned Doctor'].unique():
        doctor_patients = patient_df[patient_df['Assigned Doctor'] == doctor]
        
        # For each patient, calculate their insurance coverage and out-of-pocket costs
        for _, patient in doctor_patients.iterrows():
            disease = patient['Disease']
            insurance_company = patient['Insurance Company']
            
            # Get the insurance rate for this disease and company
            insurance_pays = 0
            if insurance_company in insurance_rates and disease in insurance_rates[insurance_company]:
                insurance_pays = insurance_rates[insurance_company][disease]
            else:
                print(f"Warning: Disease '{disease}' not found in rates for {insurance_company}")
            
            # Get doctor charge and convert to float
            try:
                doctor_charge = float(patient['Doctor Charge'])
            except (ValueError, TypeError):
                print(f"Error converting doctor charge for patient {patient['Patient ID']}")
                doctor_charge = 0
            
            # Calculate out of pocket (doctor charge - insurance pays)
            out_of_pocket = doctor_charge - insurance_pays
            
            billing_summary.append({
                'Doctor': doctor,
                'Patient ID': patient['Patient ID'],
                'Patient Name': patient['Patient Name'],
                'Disease': disease,
                'Doctor Charge': doctor_charge,
                'Insurance Company': insurance_company,
                'Insurance Pays': insurance_pays,
                'Patient Pays': out_of_pocket
            })
    
    # Create DataFrame and format currency columns
    billing_df = pd.DataFrame(billing_summary)
    currency_cols = ['Doctor Charge', 'Insurance Pays', 'Patient Pays']
    
    # Format currency columns
    for col in currency_cols:
        billing_df[col] = billing_df[col].apply(lambda x: f"${x:,.2f}" if pd.notna(x) else "")
    
    return billing_df

def df_to_docx_table(df, output_file):
    """
    Save a pandas DataFrame to a Word document as a table.
    """
    try:
        # Create a new Word document
        doc = Document()
        
        # Add a title
        doc.add_heading('Patient Disease Assignments', level=1)
        
        # Create a table
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
        
        # Add data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                # Convert numeric values to string with proper formatting
                if isinstance(value, (int, float)):
                    row_cells[i].text = f"{value:.2f}"
                else:
                    row_cells[i].text = str(value)
        
        # Save the document
        doc.save(output_file)
        return True
    except Exception as e:
        print(f"Error saving to Word document: {str(e)}")
        return False

def doctor_rates_to_dict(doctor_df):
    """
    Convert the doctor rates DataFrame to a nested dictionary:
    {
        doctor_name: {
            (disease, icd_code): rate,
            ...
        },
        ...
    }
    """
    rates_dict = {}
    for _, row in doctor_df.iterrows():
        disease = row.get('Disease Name')
        icd_code = row.get('ICD Code')
        for col in doctor_df.columns:
            if col.endswith('Rate ($)') and col not in ['Doctor A Rate ($)', 'Doctor B Rate ($)']:
                doctor_name = col.replace(' Rate ($)', '')
                if doctor_name not in rates_dict:
                    rates_dict[doctor_name] = {}
                try:
                    rate = float(row[col])
                except (ValueError, TypeError):
                    rate = None
                rates_dict[doctor_name][(disease, icd_code)] = rate
    return rates_dict
